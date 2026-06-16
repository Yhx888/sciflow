import click
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from sci_flow.config import get_data_dir, get_output_dir, ensure_directories


@click.group()
def word_cli():
    """Word 文档生成模块
    
    用于创建和管理 Word 文档，支持论文初稿生成。
    """
    ensure_directories()


@word_cli.command()
@click.option("--title", default="学术论文初稿", help="论文标题")
@click.option("--author", default="作者", help="作者姓名")
@click.option("--affiliation", default="单位", help="作者单位")
@click.option("--abstract", help="摘要内容")
def create(title, author, affiliation, abstract):
    """创建 Word 论文文档"""
    click.echo(f"📄 正在创建 Word 文档: {title}")
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author_paragraph = doc.add_paragraph(author)
    author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    affiliation_paragraph = doc.add_paragraph(affiliation)
    affiliation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading('摘要', level=1)
    if abstract:
        doc.add_paragraph(abstract)
    else:
        doc.add_paragraph('本文针对当前研究领域的问题，提出了一种新的方法...')
    
    doc.add_paragraph()
    
    doc.add_heading('关键词', level=1)
    doc.add_paragraph('关键词1；关键词2；关键词3')
    
    doc.add_paragraph()
    
    doc.add_heading('一、引言', level=1)
    doc.add_paragraph('研究背景和意义...')
    doc.add_paragraph('当前研究现状...')
    doc.add_paragraph('本文的主要贡献...')
    
    doc.add_paragraph()
    
    doc.add_heading('二、相关工作', level=1)
    doc.add_paragraph('本节介绍相关领域的研究工作...')
    
    doc.add_paragraph()
    
    doc.add_heading('三、方法', level=1)
    doc.add_paragraph('本节详细描述所提出的方法...')
    
    doc.add_paragraph()
    
    doc.add_heading('四、实验', level=1)
    doc.add_paragraph('实验设置...')
    doc.add_paragraph('实验结果与分析...')
    
    doc.add_paragraph()
    
    doc.add_heading('五、结论', level=1)
    doc.add_paragraph('本文总结...')
    doc.add_paragraph('未来工作...')
    
    doc.add_paragraph()
    
    doc.add_heading('参考文献', level=1)
    doc.add_paragraph('[1] 作者. 标题[J]. 期刊, 年份, 卷(期): 页码.')
    doc.add_paragraph('[2] 作者. 标题[C]. 会议名称, 年份: 页码.')
    
    output_path = get_output_dir() / f"{title}.docx"
    doc.save(output_path)
    
    click.echo(f"✅ Word 文档已生成: {output_path}")


@word_cli.command()
@click.option("--template", help="模板文档路径")
@click.option("--data", help="数据文件路径")
def fill(template, data):
    """从模板填充数据生成文档"""
    click.echo(f"🔄 正在从模板填充数据: {template}")
    
    if template:
        template_path = Path(template)
        if not template_path.exists():
            click.echo(f"❌ 模板文件不存在: {template}")
            return
        
        doc = Document(str(template_path))
    else:
        doc = Document()
    
    if data:
        data_path = Path(data)
        if data_path.exists():
            with open(data_path, "r", encoding="utf-8") as f:
                data_content = f.read()
            doc.add_paragraph(f"数据内容:\n{data_content}")
        else:
            click.echo(f"⚠️  数据文件不存在: {data}")
    
    output_path = get_output_dir() / "filled_document.docx"
    doc.save(output_path)
    
    click.echo(f"✅ 填充后的文档已生成: {output_path}")
