"""
SciFlow 成果生成模块
负责生成论文大纲、实验设计、调研报告、参考文献等各类学术成果
"""

from __future__ import annotations

import io
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from ..llm.client import LLMClient, get_llm_client
from .config import Config, get_config
from .literature import CitationFormatter, LiteratureManager, get_literature_manager
from .models import ExperimentDesign, Literature, Project


class ResultGenerator:
    """成果生成器类"""

    def __init__(
        self,
        config: Optional[Config] = None,
        llm: Optional[LLMClient] = None,
        literature_manager: Optional[LiteratureManager] = None,
    ):
        self.config = config or get_config()
        self.llm = llm or get_llm_client()
        self.literature_manager = literature_manager or get_literature_manager()
        self.formatter = CitationFormatter()

    def generate_outline(self, topic: str, literature: List[Literature]) -> str:
        """
        生成Markdown论文大纲

        Args:
            topic: 研究主题
            literature: 相关文献列表

        Returns:
            Markdown格式的论文大纲
        """
        current_year = datetime.now().year

        lit_summary = ""
        if literature:
            lit_summary = "## 参考文献基础\n"
            for i, lit in enumerate(literature[:5], 1):
                authors = lit.authors[0] + " 等" if len(lit.authors) > 1 else (lit.authors[0] if lit.authors else "")
                lit_summary += f"{i}. {lit.title} ({authors}, {lit.year or 'N/A'})\n"

        outline = f"""# {topic}研究报告

## 摘要

- 研究背景与问题陈述
- 研究目的与意义
- 主要研究方法
- 核心发现与贡献
- 关键词：{topic}、研究方法、实验验证

---

## 1. 引言

### 1.1 研究背景与意义
- 领域发展历程
- 现实需求与应用场景
- 理论价值与实践意义

### 1.2 研究问题与挑战
- 现有方法的局限性
- 关键科学问题
- 技术挑战分析

### 1.3 研究内容与目标
- 研究范围界定
- 具体研究目标
- 拟解决的关键问题

### 1.4 研究方法与技术路线
- 总体研究框架
- 技术路线图
- 主要方法概述

### 1.5 论文组织结构
- 各章节内容简介

---

## 2. 文献综述

### 2.1 相关研究领域概述
- 领域发展脉络
- 主要研究分支
- 研究范式演变

### 2.2 国内外研究现状
#### 2.2.1 理论研究进展
#### 2.2.2 方法技术进展
#### 2.2.3 应用实践进展

### 2.3 现有方法分析与比较
- 主流方法分类
- 各类方法优缺点
- 性能对比分析

### 2.4 研究空白与机会
- 现有研究不足
- 潜在突破方向
- 本研究切入点

---

## 3. 研究方法

### 3.1 总体框架设计
- 方法整体架构
- 核心设计思想
- 模块间关系

### 3.2 问题形式化定义
- 符号定义
- 问题建模
- 目标函数

### 3.3 核心算法/方法
#### 3.3.1 方法一：[方法名称]
- 基本原理
- 算法流程
- 关键实现细节

#### 3.3.2 方法二：[方法名称]（如适用）
- 基本原理
- 算法流程
- 关键实现细节

### 3.4 理论分析
- 收敛性分析
- 复杂度分析
- 正确性证明（如适用）

---

## 4. 实验设计与结果分析

### 4.1 实验设置
#### 4.1.1 数据集
- 数据集介绍
- 数据预处理
- 评价指标

#### 4.1.2 基线方法
- 对比方法选择
- 实现细节
- 参数设置

#### 4.1.3 实验环境
- 硬件配置
- 软件环境
- 实现框架

### 4.2 主实验结果
- 整体性能对比
- 结果表格与图表
- 统计显著性分析

### 4.3 消融实验
- 各模块贡献分析
- 参数敏感性分析
- 组件有效性验证

### 4.4 进一步分析
#### 4.4.1 定性分析
- 案例分析
- 可视化结果
- 错误分析

#### 4.4.2 效率分析
- 训练/推理时间
- 资源消耗
- 可扩展性分析

### 4.5 实验结论
- 主要发现总结
- 结果讨论

---

## 5. 讨论

### 5.1 研究发现解读
- 结果意义阐释
- 与现有研究的关系
- 理论/实践启示

### 5.2 研究局限性
- 方法局限
- 实验局限
- 适用范围

### 5.3 未来工作方向
- 方法改进方向
- 拓展研究方向
- 新的研究问题

---

## 6. 结论

- 研究工作总结
- 主要贡献回顾
- 研究意义总结

---

## 参考文献

{lit_summary}

---

## 附录

### 附录A：补充材料
### 附录B：详细证明
### 附录C：更多实验结果
"""
        return outline

    def generate_experiment_design(
        self, topic: str, literature: List[Literature]
    ) -> Dict[str, Any]:
        """
        生成实验方案

        Args:
            topic: 研究主题
            literature: 相关文献列表

        Returns:
            实验方案字典
        """
        experiment = ExperimentDesign(
            title=f"{topic}实验验证方案",
            hypothesis=f"所提出的方法在{topic}任务上能够显著优于现有基线方法，在保持效率的同时提升性能。",
            variables={
                "自变量": ["模型架构参数", "训练数据规模", "超参数设置"],
                "因变量": ["准确率/性能指标", "训练时间", "推理速度", "资源消耗"],
                "控制变量": ["随机种子", "硬件环境", "评估协议", "数据划分"],
            },
            methodology=[
                "数据预处理：按照标准流程对数据集进行清洗、标注和划分",
                "基线复现：复现3-5个当前主流的基线方法，确保公平对比",
                "模型训练：使用相同的训练策略和优化器训练所有模型",
                "评估验证：在统一的测试集上评估，使用标准评价指标",
                "统计检验：使用t检验验证结果的统计显著性",
                "消融实验：逐一验证各组件的贡献",
            ],
            expected_results=f"预期所提方法在主要评价指标上超过基线方法2-5个百分点，同时保持相当的推理效率。",
            metrics=[
                "准确率 (Accuracy)",
                "F1分数 (F1-Score)",
                "精确率 (Precision)",
                "召回率 (Recall)",
                "训练时间",
                "推理延迟",
                "GPU显存占用",
                "参数量",
            ],
            controls=[
                "所有实验使用相同的随机种子 (42)",
                "训练轮数统一设置为100轮，使用早停策略",
                "批次大小根据GPU显存统一调整",
                "学习率使用网格搜索确定最优值",
                "数据集按8:1:1划分为训练/验证/测试集",
            ],
            equipment=[
                "NVIDIA A100 GPU (或同等算力GPU)",
                "32GB以上内存",
                "500GB以上存储空间",
                "CUDA 11.8+",
                "Python 3.9+",
                "PyTorch 2.0+",
            ],
            timeline="预计8-12周完成",
        )

        return experiment.model_dump()

    def generate_report(
        self,
        topic: str,
        literature: List[Literature],
        outline: Optional[str] = None,
        experiments: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        生成完整调研报告Markdown

        Args:
            topic: 研究主题
            literature: 文献列表
            outline: 大纲（可选）
            experiments: 实验方案（可选）

        Returns:
            Markdown格式的完整报告
        """
        current_year = datetime.now().year
        date_str = datetime.now().strftime("%Y年%m月%d日")

        lit_references = self.literature_manager.generate_gbt7714(literature) if literature else ""
        lit_count = len(literature)

        report = f"""# {topic}研究调研报告

**生成日期**：{date_str}  
**文献数量**：{lit_count}篇  
**版本**：v1.0

---

## 摘要

本报告针对"{topic}"这一研究主题进行了全面的调研分析。通过系统检索和分析相关文献，梳理了该领域的研究现状、主要方法、发展趋势和存在的挑战。报告首先介绍了研究背景与意义，然后详细综述了国内外相关研究进展，对主流方法进行了分类和比较分析。在此基础上，提出了可能的研究方向和创新思路，并设计了相应的实验验证方案。本报告旨在为后续深入研究提供系统性的参考和指导。

**关键词**：{topic}；文献综述；研究方法；实验设计

---

## 1. 引言

### 1.1 研究背景与意义

随着信息技术的快速发展，{topic}近年来受到学术界和工业界的广泛关注。该领域的研究不仅具有重要的理论价值，还在实际应用中展现出巨大的潜力。从学术角度看，{topic}涉及多个学科的交叉融合，推动了相关理论和方法的创新发展。从应用角度看，相关技术已经在诸多场景中得到成功应用，产生了显著的经济和社会效益。

### 1.2 研究问题

当前{topic}领域虽然取得了长足进展，但仍面临若干关键挑战：
1. 现有方法在复杂场景下的性能仍有提升空间
2. 模型效率与精度之间的平衡尚未得到很好解决
3. 跨领域迁移和泛化能力有待加强
4. 可解释性和鲁棒性问题需要进一步研究

### 1.3 调研方法

本调研采用系统性文献综述方法，按照以下步骤进行：
1. 确定检索关键词和数据库
2. 筛选和评估相关文献
3. 数据提取与整理
4. 定性与定量分析
5. 综合归纳与总结

---

## 2. 文献综述

### 2.1 领域发展概况

{topic}的研究可以追溯到早期的相关工作。经过多年发展，该领域经历了从传统方法到深度学习方法的转变，性能得到了显著提升。特别是近年来，随着大模型技术的突破，{topic}领域迎来了新的发展机遇。

### 2.2 主要研究方向

根据对检索到的{lit_count}篇文献的分析，当前{topic}的研究主要集中在以下几个方向：

"""
        if literature:
            report += """#### 2.2.1 基于深度学习的方法

深度学习方法在{}任务上取得了突破性进展。研究者们提出了多种网络架构和训练策略，显著提升了模型性能。代表性工作包括：

""".format(topic)
            for i, lit in enumerate(literature[:3], 1):
                authors = lit.authors[0] + (" 等" if len(lit.authors) > 1 else "")
                report += f"- [{i}] {authors}({lit.year})提出了{lit.title[:30]}...，该方法...\n"

            report += """
#### 2.2.2 轻量化与效率优化方法

为解决模型部署问题，许多研究关注模型压缩和加速。主要技术包括知识蒸馏、模型剪枝、量化等。

#### 2.2.3 跨模态与多任务学习

利用多模态信息和多任务学习来提升模型性能和泛化能力是当前的研究热点之一。

"""

        report += """### 2.3 方法比较分析

| 方法类别 | 代表方法 | 优点 | 缺点 | 适用场景 |
|---------|---------|------|------|---------|
| 传统方法 | 统计方法、规则系统 | 可解释性强、计算简单 | 性能有限、特征工程复杂 | 简单场景、小数据 |
| 深度学习 | CNN、RNN、Transformer | 性能强大、端到端学习 | 数据需求大、计算成本高 | 复杂场景、大数据 |
| 轻量化方法 | 蒸馏、剪枝、量化 | 部署友好、效率高 | 精度损失、设计复杂 | 边缘设备、实时应用 |

### 2.4 研究趋势分析

通过文献分析，可以观察到以下发展趋势：
1. **模型规模持续增大**：大模型成为主流，但小模型优化同样重要
2. **多模态融合**：文本、图像、语音等多模态联合处理
3. **高效学习**：小样本、零样本学习成为重要方向
4. **可解释AI**：模型可解释性越来越受重视
5. **实际部署**：从实验室走向真实应用场景

---

## 3. 研究思路与建议

### 3.1 潜在创新点

基于文献分析，以下方向可能具有创新潜力：

1. **架构创新**：设计更适合{topic}任务的新型网络架构
2. **训练策略**：提出更高效的训练方法和优化策略
3. **数据利用**：探索小样本、弱监督等数据高效利用方式
4. **融合方法**：将多种技术有机结合，取长补短

### 3.2 推荐研究方向

综合考虑创新性、可行性和应用价值，建议优先考虑以下方向：

**研究方向：面向{topic}的高效方法研究**
- 核心思想：在保持高性能的同时显著提升模型效率
- 技术路线：架构改进 + 训练优化 + 知识蒸馏
- 预期贡献：提出新的高效模型，在精度和速度上取得更好平衡
- 可行性：⭐⭐⭐⭐⭐（技术积累充分，实验条件成熟）

---

## 4. 实验方案设计

"""
        if experiments:
            report += f"""### 4.1 实验目标

{experiments.get('hypothesis', '验证所提方法的有效性')}

### 4.2 实验设置

**评价指标**：
"""
            for metric in experiments.get('metrics', []):
                report += f"- {metric}\n"

            report += """
**控制变量**：
"""
            for control in experiments.get('controls', []):
                report += f"- {control}\n"

            report += """
### 4.3 实验流程

"""
            for i, step in enumerate(experiments.get('methodology', []), 1):
                report += f"{i}. {step}\n"

            report += f"""
### 4.4 预期结果

{experiments.get('expected_results', '')}

**时间安排**：{experiments.get('timeline', '预计8-12周')}
"""
        else:
            report += """实验方案详见后续详细设计。"""

        report += f"""

---

## 5. 结论

本报告对{topic}领域进行了系统性的文献调研和分析。通过梳理{lit_count}篇相关文献，总结了该领域的研究现状、主要方法和发展趋势。分析表明，{topic}领域虽然已经取得了显著进展，但仍存在若干值得深入研究的问题。建议后续研究可以从高效方法设计、多模态融合、小样本学习等方向切入，开展创新性研究工作。

---

## 参考文献

{lit_references}

---

*本报告由SciFlow AI科研助手自动生成，仅供参考*
"""
        return report

    def generate_bibtex(self, literature: List[Literature]) -> str:
        """
        生成BibTeX文件内容

        Args:
            literature: 文献列表

        Returns:
            BibTeX格式的参考文献
        """
        return self.literature_manager.generate_bibtex(literature)

    def generate_word(
        self,
        report_content: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        生成Word文档（使用python-docx，返回文件路径）

        Args:
            report_content: Markdown报告内容
            metadata: 文档元数据

        Returns:
            Word文档路径
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return "python-docx库未安装，无法生成Word文档。请运行: pip install python-docx"

        exports_dir = self.config.get_exports_dir()
        word_path = exports_dir / "report.docx"

        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)

        title = doc.add_heading(metadata.get("title", "研究报告") if metadata else "研究报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if metadata:
            if "author" in metadata:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"作者：{metadata['author']}")
            if "affiliation" in metadata:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"单位：{metadata['affiliation']}")
            if "date" in metadata:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(f"日期：{metadata['date']}")

        doc.add_page_break()

        lines = report_content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("#### "):
                doc.add_heading(line[5:], level=4)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
            elif line.startswith("---"):
                doc.add_paragraph("─" * 50)
            else:
                clean_line = line.replace("**", "").replace("*", "")
                doc.add_paragraph(clean_line)

        doc.save(str(word_path))
        return str(word_path)

    def generate_project_bundle(self, project: Project) -> bytes:
        """
        打包所有成果为ZIP文件

        Args:
            project: 项目对象

        Returns:
            ZIP文件字节内容
        """
        buf = io.BytesIO()

        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            readme_content = f"""# {project.topic} - SciFlow科研项目

**作者**: {project.author or '未指定'}
**单位**: {project.affiliation or '未指定'}
**创建时间**: {project.created_at.strftime('%Y-%m-%d') if hasattr(project.created_at, 'strftime') else project.created_at}
**项目描述**: {project.description}

## 文件说明

- `report.md`: 完整调研报告（Markdown格式）
- `references.bib`: BibTeX参考文献
- `outline.md`: 论文大纲
- `project.json`: 项目元数据

*由SciFlow AI科研助手生成*
"""
            zf.writestr("README.md", readme_content)

            if project.report:
                zf.writestr("report.md", project.report)

            project_literature = self.literature_manager.get_project_literature(project.id)
            if project_literature:
                bibtex_content = self.generate_bibtex(project_literature)
                zf.writestr("references.bib", bibtex_content)

                gbt_content = self.literature_manager.generate_gbt7714(project_literature)
                zf.writestr("references_gbt7714.txt", gbt_content)

                apa_content = self.literature_manager.generate_apa(project_literature)
                zf.writestr("references_apa.txt", apa_content)

            if project.outline:
                zf.writestr("outline.md", project.outline)

            import json
            project_data = {
                "id": project.id,
                "topic": project.topic,
                "description": project.description,
                "author": project.author,
                "affiliation": project.affiliation,
                "status": project.status.value if hasattr(project.status, 'value') else str(project.status),
                "created_at": str(project.created_at),
                "literature_count": len(project.literature),
            }
            zf.writestr("project.json", json.dumps(project_data, ensure_ascii=False, indent=2))

        buf.seek(0)
        return buf.getvalue()


_generator_instance: Optional[ResultGenerator] = None


def get_result_generator(
    config: Optional[Config] = None,
    llm: Optional[LLMClient] = None,
    literature_manager: Optional[LiteratureManager] = None,
) -> ResultGenerator:
    """获取成果生成器单例"""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = ResultGenerator(config, llm, literature_manager)
    return _generator_instance


def reset_result_generator() -> None:
    """重置成果生成器（主要用于测试）"""
    global _generator_instance
    _generator_instance = None
